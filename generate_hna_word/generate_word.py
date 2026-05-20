import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
import win32com.client
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime
from numbers import Number
import os
import glob
import re
import shutil
import tempfile


script_dir = os.getcwd()

folder_path =  os.path.join(script_dir,'input')
all_files = [f for f in os.listdir(folder_path) if not f.startswith('HNA Municipal Template')
             and f.lower().endswith('.xlsx')
             ]
output_folder = os.path.join(script_dir, 'output')
word_path = os.path.join(script_dir, 'input', "HNA Municipal Template 2026.05.04.docx")

for filename in all_files:
    


    temp_path = tempfile.mktemp(suffix='.docx')
    shutil.copy(word_path, temp_path)

    excel_path = os.path.join(folder_path, filename)
    excel_file = excel_path
    word_file = temp_path
    doc = Document(word_file)
    

    excel = win32com.client.Dispatch("Excel.Application")
    excel.Visible = False
    wb = excel.Workbooks.Open(excel_path)
    wb.Save()

    wb.Close(SaveChanges=True)


    #### Sheet: pop_now ####
    sheet = 'popn_now'
    columns = ['B:J',
            'B:E',
            'B:E'
            ]
    header_num = [
        50,
        160,
        191
    ]
    rows_num = [
        8,
        5,
        10
    ]
    titles = [
        'Population estimates by age group (2025) compared to most recent census (2021) and comparison of age distribution by geography',
        'Five-year mobility by origin and geography, 2021',
        'Key population group shares of total population by geography, 2021'
    ]
    for i in range(len(columns)):
        df = pd.read_excel(excel_file, sheet_name = sheet, usecols = columns[i], header = header_num[i], nrows = rows_num[i])
        table_par_idx = None
        search_text = titles[i]
        for j, para in enumerate(doc.paragraphs): # Finds paragraph with the title and gets index
            if search_text in para.text:
                table_par_idx = j
                break
        table_to_replace = None
        if table_par_idx is not None:
            par_elem = doc.paragraphs[table_par_idx]._element
            for tbl in doc.tables: # Gets the table after the title
                if tbl._element.getprevious() is par_elem:
                    table_to_replace = tbl
                    break
        if table_to_replace: # Fills each table according to the title
            if search_text == titles[0]:
                for a in [2, 3, 5, 6, 7]:
                    for b in range(1, 9):
                        cell = table_to_replace.cell(a, b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val:,.0f}' if a == 2 and isinstance(val, (float, int))
                            else f'{val*100:,.0f}%' if isinstance(val,(float, int))
                            else str(val)
                            )

                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[1]:
                for a in range(min(len(table_to_replace.rows), len(df))):
                    for b in range(min(len(table_to_replace.rows[a].cells), len(df.columns))):
                        cell = table_to_replace.cell(a, b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            "" if pd.isna(val) or val in [None, '']
                            else f'{val*100:.1f}%' if isinstance(val, (float, int)) 
                            else str(val))
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[2]:
                for a in range(min(len(table_to_replace.rows), len(df))):
                    for b in range(min(len(table_to_replace.rows[a].cells), len(df.columns))):
                        cell = table_to_replace.cell(a, b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None,'']
                            else f'{val*100:.1f}%' if isinstance(val, (float, int)) 
                            else str(val))
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold



    #### Sheet: hh ####
    sheet = 'hh'
    columns = ['B:H',
            'B:G'
            ]
    header_num = [
        35,
        84
    ]
    rows_num = [
        25,
        11
    ]
    titles = [
        'Summary of household characteristics by tenure and geography, 2021',
        'Estimated household suppression by age of primary household maintainer, 2021'
    ]

    for i in range(len(columns)):
        df = pd.read_excel(excel_file, sheet_name = sheet, usecols = columns[i], header = header_num[i], nrows = rows_num[i])
        table_par_idx = None
        search_text = titles[i]
        for j, para in enumerate(doc.paragraphs):
            if search_text in para.text:
                table_par_idx = j
                break
        table_to_replace = None
        if table_par_idx is not None:
            par_elem = doc.paragraphs[table_par_idx]._element
            for tbl in doc.tables:
                if tbl._element.getprevious() is par_elem:
                    table_to_replace = tbl
                    break
        if table_to_replace:
            if search_text == titles[0]:
                rows = [2] + list(range(4,12)) + list(range(13,18)) + list(range(19,25))
                for a in rows : # Rows
                    for b in range(1,8): # Columns
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a, b-1]
                        cell.text = (
                            '' if pd.isna(val) or val in [None,'']
                            else f"{val:,.0f}" if b == 2 and a !=24 and isinstance(val, (float, int))
                            else f"{val:,.0f}" if isinstance(val, (float, int)) and b in range(5,8) and a==2
                            else f"{val:,.1f}" if isinstance(val, (float, int)) and a ==24
                            else f"{val*100:,.0f}%" if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[1]:
                for a in range(2, 11): # Rows
                    for b in range(1, 6): # Columns
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc [a, b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None,'']
                            else f"{val:.2f}" if isinstance(val, (float, int)) and b in range(1,3)
                            else f"{val:,.0f}" if isinstance(val, (float, int)) and b in range(3,6)
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold

    #### Sheet: econ ####

    sheet = 'econ'
    columns = ['B:E',
            'B:E',
            'B:H',
            'I:M'
            ]
    header_num = [
        21, 60, 97, 126
    ]
    rows_num = [
        4, 11, 7, 10
    ]
    titles = [
        'Labour force metrics by geography, 2021',
        'Top 10 local NAICS industries of employment compared to higher level geographies, 2021',
        'Top 5 origins and destinations of commuters relative to the local community, 2021',
        'Share of population within select transit service areas* by geography, 2024'
        ]

    for i in range(len(columns)):
        df = pd.read_excel(excel_file, sheet_name = sheet, usecols = columns[i], header = header_num[i], nrows = rows_num[i])
        table_par_idx = None
        search_text = titles[i]
        for j, para in enumerate(doc.paragraphs):
            if search_text in para.text:
                table_par_idx = j
                break
        table_to_replace = None
        if table_par_idx is not None:
            par_elem = doc.paragraphs[table_par_idx]._element
            for tbl in doc.tables:
                if tbl._element.getprevious() is par_elem:
                    table_to_replace = tbl
                    break
        if table_to_replace:
            if search_text == titles[0]:
                for a in range(min(len(table_to_replace.rows), len(df))):
                    for b in range(min(len(table_to_replace.rows[a].cells), len(df.columns))):
                        cell = table_to_replace.cell(a, b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment                        
                        val = df.iloc[a, b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f"{val*100:,.1f}%" if isinstance(val, (float, int)) 
                            else str(val))
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[1]:
                for a in range(min(len(table_to_replace.rows), len(df))):
                    for b in range(min(len(table_to_replace.rows[a].cells), len(df.columns))):
                        cell = table_to_replace.cell(a, b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment                        
                        val = df.iloc[a, b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val*100:,.1f}%' if isinstance(val, (float, int)) 
                            else str(val))
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[2]:
                for a in range(1,7): # rows
                    for b in range(0,3): # Columns
                        cell = table_to_replace.cell(a, b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f"{val}" if b == 0 and isinstance(val,(float,int))
                            else f"{val:,.0f}" if b ==1 and isinstance(val, (float, int))
                            else f'{val*100:,.0f}%'if b ==2 and isinstance(val,(float, int))
                            else str(val)
                        )
                        
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
                    for b in range(4, 7):
                        cell = table_to_replace.cell(a, b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a, b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else str(val) if b ==4
                            else f"{val:,.0f}" if isinstance(val, (float, int)) and b ==5
                            else f'{val*100:,.0f}%' if isinstance(val, (float, int)) and b ==6
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[3]:
                for a in range(2, 10):
                    for b in range(1,5):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bgold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val*100:,.1f}%' if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
                

    #### Sheet: income ####


    sheet = 'income'
    columns = ['B:E',
            'B:H',
            'B:H',
            'B:E',
            'B:G',
            'I:O'
            ]
    header_num = [
        37, 93, 182,139, 265, 265
    ]
    rows_num = [
        16, 12, 7,16, 3, 6
    ]
    titles = [
        'Median household income by household characteristic and geography, 2021 Census',
        'Household income categories by household characteristic, 2021 Census',
        'Select government transfer incomes by geography, 2021 Census',
        'Percent of persons in poverty (Market Basket Measure) by characteristic and geography, 2021 Census',
        'Estimated poverty line and living wage income thresholds, 2024',
        'Estimated "ALICE" households by household size based on poverty line and living wage income thresholds, 2021'
    ]


    for i in range(len(columns)):
        df = pd.read_excel(excel_file, sheet_name = sheet, usecols = columns[i], header = header_num[i], nrows = rows_num[i])
        table_par_idx = None
        search_text = titles[i]
        for j, para in enumerate(doc.paragraphs):
            if search_text in para.text:
                table_par_idx = j
                break
        table_to_replace = None
        if table_par_idx is not None:
            par_elem = doc.paragraphs[table_par_idx]._element
            for tbl in doc.tables:
                if tbl._element.getprevious() is par_elem:
                    table_to_replace = tbl
                    break
        if table_to_replace:
            if search_text == titles[0]:
                rows = [1] + list(range(3,5)) + list(range(6,11)) + list(range(12,16))
                for a in rows:
                    for b in range(1,4):
                        cell = table_to_replace.cell(a, b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = f'${val:,.0f}' if isinstance(val, (float, int)) else str(val)
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[1]:
                for a in range(2,12):
                    for b in range(2,8):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = (df.iloc[2, b-1] if a in [2,3] and b in range(3,8) else df.iloc[a,b-1])
                        cell.text = (
                            'Total' if a in [1,2,3] and b ==2
                            else '' if pd.isna(val) or val in [None,'']
                            else f'{val:,.0f}' if isinstance(val, (float, int)) and a in range(4,12) and b ==2
                            else f'{val*100:,.0f}%' if isinstance(val, (float, int)) and a in range(4,12) and b in range(3,8)
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text ==titles[2]:
                for a in range(2,7):
                    for b in range(1,7):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f"{val:,.0f}" if isinstance(val, (float, int)) and b in [1,3,5]
                            else f"{val*100:,.0f}%" if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[3]:
                rown = [1] + list(range(3,9)) + list(range(10,16))
                for a in rown:
                    for b in range(2,5):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = (
                            df.iloc[a,2] if b in [2,3]
                            #else df.iloc[a,1] if b ==2
                            else df.iloc[a,3] if b == 4
                            else df.iloc[a,b]
                            )
                        cell.text = (
                                    '' if pd.isna(val) or val in [None, '']
                            else f"{val*100:,.0f}%" if isinstance(val, (float, int)) and b in [2,3]
                            else f"{val*100:,.0f}%" if isinstance(val, (float, int)) and b == 2
                            else f"{val*100:,.0f}%" if isinstance(val, (float, int)) and b == 4
                            else str(val)
                                    )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[4]:
                for a in [1,2]:
                    for b in range(1,6):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'${val:,.0f}' if isinstance(val,(float,int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[5]:
                for a in range(1, 6):
                    for b in range(1,7):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val*100:,.0f}%' if a == 5 and isinstance(val, (float, int))
                            else f'{val:,.0f}' if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold






    #### Sheet: dwell ####
    sheet = 'dwell'
    columns = ['B:H',
            'B:E',
            'G:J',
            'B:E',
            'G:J',
            'B:E',
            'B:H',
            'B:H',
            'B:K',
            'B:H',
            'B:E',
            'B:D',
            'F:H',
            'B:H',
            'B:H'
            ]
    header_num = [
        98,
        174,
        174,
        183,
        183,
        140,
        252,
        282,
        62,
        30,
        218,
        329,
        329,
        290,
        370
    ]
    rows_num = [
        11,
        6,
        6,
        5,
        5,
        4,
        5,
        5,
        5,
        8,
        4,
        19,
        6,
        5,
        10
    ]
    titles = [
        'Dwellings occupied by a usual resident by structure type, tenure, and geography, 2021',
        'Rental units afforded by very-low and low-income households by geography, 2021',
        'Rental units afforded by very-low and low-income households built between 2016 and 2021 by geography',
        'Rental units afforded by very-low and low-income households lost from 2016 to 2021 by geography',
        'Net change in rental units (2016 to 2021) afforded by very low- and low-income households by geography',
        'Subsidized rental housing by geography, 2021',
        'Short-term rental (STR) and potential long-term dwellings (PLTDs) by geography',
        'Estimated share of buildings, dwelling units, and people subject to moderate and high susceptibility to flooding',
        "Total and share of age of construction for dwellings occupied by usual residents, 2021",
        'Recent historical permits by structure type and percent of total dwellings',
        'Private market rental units that are priced below-market, 2021',
        'Select local climate estimates and national percentiles based on Canadian climate data and scenarios',
        'Key Census metrics pertaining to housing and demographic vulnerability to climate conditions, 2021',
        'Estimated total and share of buildings, dwelling units, and people subject to coastal flooding based on a 1:100 year flood scenario, 2021',
        'Summary of non-market housing by type of housing and geography, Spring 2026'
    ]


    for i in range(len(columns)):
        df = pd.read_excel(excel_file, sheet_name = sheet, usecols = columns[i], header = header_num[i], nrows = rows_num[i])
        table_par_idx = None
        search_text = titles[i]
        for j, para in enumerate(doc.paragraphs):
            if search_text in para.text:
                table_par_idx = j
                break
        table_to_replace = None
        if table_par_idx is not None:
            par_elem = doc.paragraphs[table_par_idx]._element
            for tbl in doc.tables:
                if tbl._element.getprevious() is par_elem:
                    table_to_replace = tbl
                    break
        if table_to_replace:
            if search_text == titles[0]:
                for a in range(2,11):
                    for b in range(0,8):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a, b-1]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f"{val:,.0f}" if isinstance(val, (float, int)) and (b == 2 or (a==2 and b in range(5,8)))
                            else f"{val*100:,.0f}%" if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[1]:
                for a in range(1,6):
                    for b in range(1, 4):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f"{val*100:,.0f}%" if isinstance(val, (float, int)) and a == 5
                            else f"{val:,.0f}" if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text ==titles[2]:
                for a in range(1,6):
                    for b in range(1,4):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f"{val*100:,.0f}%" if isinstance(val, (float, int)) and a == 5
                            else f"{val:,.0f}" if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[3]:
                for a in range(1,5):
                    for b in range(1,4):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val:,.0f}' if isinstance(val, (float, int)) 
                            else str(val))
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[4]:
                for a in range(1,5):
                    for b in range(1,4):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val:,.0f}' if isinstance(val, (float, int))
                                    else str(val))
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[5]:
                for a in range(1, 4):
                    for b in range(1, 4):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val*100:,.0f}%' if isinstance(val, (float, int)) and a ==3
                            else f'{val:,.0f}' if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[6]:
                for a in range(2, 5):
                    for b in range(1,7):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val:,.0f}' if isinstance(val, (float, int)) and a ==2
                            else f'{val*100:,.0f}%' if isinstance(val, (float,int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[7]:
                for a in range(2,5):
                    for b in range(1,7):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val*100:,.0f}%' if isinstance(val, (float,int))
                            else str(val))
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[8]:
                for a in range(1,5):
                    for b in range(1,10):
                        cell =  table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None,'']
                            else f'{val*100:,.0f}%' if isinstance(val, (float,int)) and a in [2,3,4] and b in list(range(2,10))
                            else f'{val:,.0f}' if isinstance(val, (float, int)) and a == 1
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[9]:
                for a in range(1,8):
                    for b in range(1,7):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val*100:,.1f}%' if isinstance(val, (float, int)) and a == 7
                            else f'{val:,.0f}' if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[10]:
                for a in range(1,4):
                    for b in range(1,4):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val*100:,.0f}%' if isinstance(val, (float, int)) and a == 3
                            else f'{val:,.0f}' if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[11]:
                for a in range(2,19):
                    for b in range(1,3):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            "" if pd.isna(val) or val in [None, '']
                            else f'{val:,.0f}' if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[12]:
                for a in range(1,6):
                    for b in range(1,3):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val*100:,.0f}%' if isinstance(val, (float, int)) and b == 1
                            else f'{val:,.0f}' if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[13]:
                for a in range(2,5):
                    for b in range(1,7):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val*100:,.1f}%' if isinstance(val, (float, int)) and b in [2,4,6]
                            else f'{val:,.0f}' if isinstance(val, (float, int)) and b in [1,3,5]
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[14]:
                for a in range(2,10):
                    for b in range(1,7):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val:,.0f}' if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold








    #### Sheet: rental ####


    sheet = 'rental'
    columns = ['I:M',
            'B:F',
            'B:H']
    header_num = [62,
                30, 88]
    rows_num = [7,
                11, 6]
    titles = ['Primary and secondary rental markets by dwelling size share and geography, 2021',
            'Change in average rents versus change in provincial salaries and wages (2020 to 2025), and 2025 vacancy rate, by dwelling size and geography',
            'Estimated affordable rent thresholds, 2026'
            ]



    for i in range(len(columns)):
        df = pd.read_excel(excel_file, sheet_name = sheet, usecols = columns[i], header = header_num[i], nrows = rows_num[i])
        table_par_idx = None
        search_text = titles[i]
        for j, para in enumerate(doc.paragraphs):
            if search_text in para.text:
                table_par_idx = j
                break
        table_to_replace = None
        if table_par_idx is not None:
            par_elem = doc.paragraphs[table_par_idx]._element
            for tbl in doc.tables:
                if tbl._element.getprevious() is par_elem:
                    table_to_replace = tbl
                    break
        if table_to_replace:
            if search_text == titles[0]:
                for a in range(2,7):
                    for b in range(1,5):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f"{val:,.0f}" if isinstance(val, (float, int)) and a == 2
                            else f"{val*100:,.0f}%" if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[1]:
                for a in range(2,7):
                    for b in range(1,5):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text =( 
                            '' if pd.isna(val) and val in [None, '']
                            else f"{val*100:,.1f}%" if isinstance(val, (float, int)) 
                            else str(val))
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
                for a in range(9,11):
                    for b in [1]:
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text =(
                            '' if pd.isna(val) or val in [None, '']
                            else f"{val*100:,.1f}%" if isinstance(val, (float, int)) else str(val))
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[2]:
                for a in range(2,6):
                    for b in range(2,7):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'${val:,.0f}' if isinstance(val, (float, int)) else str(val))
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold



    #### Sheet: sales ####


    sheet = 'sales'
    columns = ['B:H']
    header_num = [30]
    rows_num = [11]
    titles = ['Change in average house prices versus change in average salaries and wages (2020 to 2025) by dwelling size and geography',
            ]


    for i in range(len(columns)):
        df = pd.read_excel(excel_file, sheet_name = sheet, usecols = columns[i], header = header_num[i], nrows = rows_num[i])
        table_par_idx = None
        search_text = titles[i]
        for j, para in enumerate(doc.paragraphs):
            if search_text in para.text:
                table_par_idx = j
                break
        table_to_replace = None
        if table_par_idx is not None:
            par_elem = doc.paragraphs[table_par_idx]._element
            for tbl in doc.tables:
                if tbl._element.getprevious() is par_elem:
                    table_to_replace = tbl
                    break
        if table_to_replace:
            if search_text == titles[0]:
                for a in range(2, 7):
                    for b in range(1, 7):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f"{val*100:,.0f}%" if isinstance(val, (float, int)) and b in [2,4,6]
                            else f"${val:,.0f}" if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
                for a in range(9,11):
                    for b in [1]:
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val*100:,.0f}%' if isinstance(val, (float, int)) 
                            else str(val))
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold

    #### Sheet: need ####
    sheet = 'need'
    columns = ['B:E',
            'B:I',
            'H:K',
            'B:E'
            ]
    header_num = [103,
                196,
                168,
                220

    ]
    rows_num = [8,
                7,
                20,
                4

    ]

    titles = [
        'Summary of local housing need indicators, 2021',
        'Households in core housing need by household size and income category, 2021',
        'Rate of core housing need by vulnerable population group and geography, 2021',
        'Share of household income allocated to housing and transportation by geography, 2024'
    ]



    for i in range(len(columns)):
        df = pd.read_excel(excel_file, sheet_name = sheet, usecols = columns[i], header = header_num[i], nrows = rows_num[i])
        table_par_idx = None
        search_text = titles[i]
        for j, para in enumerate(doc.paragraphs):
            if search_text in para.text:
                table_par_idx = j
                break
        table_to_replace = None
        if table_par_idx is not None:
            par_elem = doc.paragraphs[table_par_idx]._element
            for tbl in doc.tables:
                if tbl._element.getprevious() is par_elem:
                    table_to_replace = tbl
                    break
        if table_to_replace:
            if search_text == titles[0]:
                for a in range(1,8):
                    for b in range(1,4):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val:,.0f}' if isinstance(val, (float, int)) and a == 1
                            else f'{val*100:,.1f}%' if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[1]:
                for a in range(1,7):
                    for b in range(1,8):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val or val in [None, ''])
                            else f'${val:,.0f}' if isinstance(val, (float, int)) and b == 1
                            else f'{val:,.0f}' if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[2]:
                for a in range(1,20):
                    for b in range(1,4):
                        cell = table_to_replace.cell(a, b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val*100:,.1f}%' if isinstance(val, (float, int))
                            else str(val))
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[3]:
                for a in range(1,4):
                    for b in range(1,4):
                        cell = table_to_replace.cell(a, b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val*100:,.1f}%' if isinstance(val, (float, int)) 
                            else str(val))
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold


    #### Sheet: pop_fut ####
    sheet = 'popn_fut'
    columns = ['B:J']
    header_num = [103]
    rows_num = [7]

    titles = [
        'Anticipated population by age group over next 10 years, medium growth scenario'
    ]

    for i in range(len(columns)):
        df = pd.read_excel(excel_file, sheet_name = sheet, usecols = columns[i], header = header_num[i], nrows = rows_num[i])
        table_par_idx = None
        search_text = titles[i]
        for j, para in enumerate(doc.paragraphs):
            if search_text in para.text:
                table_par_idx = j
                break
        table_to_replace = None
        if table_par_idx is not None:
            par_elem = doc.paragraphs[table_par_idx]._element
            for tbl in doc.tables:
                if tbl._element.getprevious() is par_elem:
                    table_to_replace = tbl
                    break
        if table_to_replace:
            if search_text == titles[0]:
                for a in [2,3,5,6]:
                    for b in range(1,9):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val:,.0f}' if isinstance(val, (float, int)) and a == 2
                            else f'{val*100:,.0f}%' if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold

    #### Sheet: calc ####
    sheet = 'projection'
    columns =['B:H', 'R:X', 'B:E', 'B:D']
    header_num = [269, 322, 455, 525]
    rows_num = [7,7, 4, 3]

    titles = [
        'Anticipated new dwellings demand by dwelling size, over next 10 years',
        'Anticipated dwelling demand by household size and income category, by 2035',
        'Net change to demand based on the integration of communities tied economically (via employment and affordability), 2021 estimates',
        'Potential demand for senior housing'
    ]

    for i in range(len(columns)):
        df = pd.read_excel(excel_file, sheet_name = sheet, usecols = columns[i], header=header_num[i], nrows = rows_num[i])
        table_par_idx = None
        search_text = titles[i]
        for j, para in enumerate(doc.paragraphs):
            if search_text in para.text:
                table_par_idx = j
                break
        table_to_replace = None
        if table_par_idx is not None:
            par_elem = doc.paragraphs[table_par_idx]._element
            for tbl in doc.tables:
                if tbl._element.getprevious() is par_elem:
                    table_to_replace = tbl
                    break
        if table_to_replace:
            if search_text == titles[0]:
                for a in range(2,7):
                    for b in range(1,7):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val:,.0f}' if isinstance(val, (float, int)) and b in [1,3,5]
                            else f'{val*100:,.0f}%' if isinstance(val,(float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[1]:
                for a in range(1,7):
                    for b in range(1, 7):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            "" if pd.isna(val) or val in [None, ""] 
                            else f'{val:,.0f}' if isinstance(val, (float, int)) 
                            else str(val))
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[2]:
                for a in range(1,4):
                    for b in range(1,4):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val:,.0f}' if isinstance(val, (float, int)) 
                            else str(val))
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[3]:
                for a in [1,2]:
                    for b in [1,2]:
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else (f'{val:,.0f}' if isinstance(val, Number) 
                                    else str(val))
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold

    #### Sheet: vuln ####

    sheet = 'vuln'
    columns =['B:E', 'B:G']
    header_num = [33, 46]
    rows_num = [10, 5]

    titles = ['Summary of vulnerability indices by measure and geographic comparison, 2021',
            'VALIDATED VULNERABILITY INDICES'
            ]

    for i in range(len(columns)):
        df = pd.read_excel(excel_file, sheet_name = sheet, usecols = columns[i], header = header_num[i], nrows = rows_num[i])
        table_par_idx = None
        search_text = titles[i]
        matches = [j for j, para in enumerate(doc.paragraphs) if search_text in para.text]
        table_par_idx = matches[0]
        table_to_replace=None
        if table_par_idx is not None:
            par_elem = doc.paragraphs[table_par_idx]._element
            target_idx=1 if titles[i] in ['VALIDATED VULNERABILITY INDICES'] else 0
            tbl_count =0
            sibling =  par_elem.getnext()
            while sibling is not None:
                if sibling.tag == qn('w:tbl'):
                    if tbl_count == target_idx:
                        for tbl in doc.tables:
                            if tbl._element is sibling:
                                table_to_replace = tbl
                                break
                        break
                    tbl_count+=1
                sibling = sibling.getnext()
        if table_to_replace:
            if search_text == titles [0]:
                for a in range(1,10):
                    for b in range(2,4):
                        cell = table_to_replace.cell(a, b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text =( 
                            '' if pd.isna(val) or val in [None, '']
                            else str(val))
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles [1]:
                for a in range(1,5):
                    for b in [1,3,5]:
                        cell = table_to_replace.cell(a, b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text =( 
                            '' if pd.isna(val) or val in [None, '']
                            else str(val))
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold



    #### Sheet: vuln_profile ####

    sheet = 'vuln_profile'
    columns =['B:G']
    header_num = [0]
    rows_num = [24]

    titles = ['KEY VULNERABILITY INDICATORS'
            ]

    for i in range(len(columns)):
        df = pd.read_excel(excel_file, sheet_name = sheet, usecols = columns[i], header = header_num[i], nrows = rows_num[i])
        table_par_idx = None
        search_text = titles[i]
        for j, para in enumerate(doc.paragraphs):
            if search_text in para.text:
                table_par_idx = j
                break
        table_to_replace = None
        if table_par_idx is not None:
            par_elem = doc.paragraphs[table_par_idx]._element
            sibling = par_elem.getnext()
            while sibling is not None:
                if sibling.tag ==qn('w:tbl'):
                    for tbl in doc.tables:
                        if tbl._element is sibling:
                            table_to_replace = tbl
                            break
                    break
                sibling = sibling.getnext()
        if table_to_replace:
            if search_text == titles [0]:
                for a in range(2, 24):
                    for b in range(2,5):
                        cell = table_to_replace.cell(a, b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text =( 
                            '' if pd.isna(val) or val in [None, '']
                            else f'${val:,.0f}' if isinstance(val, (float, int)) and a == 12
                            else f'{val:,.0f}' if isinstance(val, (float, int)) and a in [19,20]
                            else f'{val*100:,.1f}%' if isinstance(val, (float, int))
                            else str(val))
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold




    #### Sheet: federal ####

    sheet = 'federal_req'
    columns =['B:E', 'B:E', 'B:E', 'B:E',
            'G:J', 'G:J', 'G:J', 'G:J', 'G:J',
            'L:P', 'L:S', 'L:O', 'L:O', 'L:R']
    header_num = [1, 17, 26, 59, 1, 35, 43, 70, 89, 1, 10, 20, 35, 65]
    rows_num = [13, 6, 30, 30, 31, 5, 24, 16, 40, 6, 7, 12,27, 7]

    titles = ['Table 2.2.1: Population',
            'Table 2.2.2: Demographic Information',
            'Table 3.1.1: Household Income and Profile',
            'Table 3.4.1: Economy and Labour Force',
            'Table 5.2.1: Housing Units Currently Occupied / Available',
            'Table 5.7.1: Current Non-Market Housing Units',
            'Table 5.9.1: Housing Values',
            'Table 5.9.2: Change in Housing Stock ',
            'Table 3.6.4: Households in Core Housing Need',
            'Income Categories and Affordable Shelter Costs (HART)',
            'Core Housing Need by Income Category and Household Size',
            'Table 6.3.1: Anticipated Population by 2035',
            'Table 6.3.2: Anticipated Households by 2035',
            'Table 6.1.1: Projected Households by Household Size and Income Category'

            ]

    for i in range(len(columns)):
        df = pd.read_excel(excel_file, sheet_name=sheet, usecols=columns[i], header=header_num[i], nrows=rows_num[i])
        search_text = titles[i]

        # Walk body XML children to find title paragraph, then the next table
        table_to_replace = None
        found_title = False
        body = doc.element.body
        for child in body:
            if not found_title:
                if child.tag == qn('w:p'):
                    text = ''.join(t.text or '' for t in child.iter(qn('w:t')))
                    if search_text in text:
                        found_title = True
            else:
                if child.tag == qn('w:tbl'):
                    for tbl in doc.tables:
                        if tbl._element is child:
                            table_to_replace = tbl
                            break
                    break  # stop after finding first table after title

        if table_to_replace: # Fills each table according to the title
            if search_text == titles[0]:
                cell1 = table_to_replace.cell(1,1)
                for a in range(1,13):
                    for b in [2,3]:
                        cell = table_to_replace.cell(a, b)
                        orig_run = cell1.paragraphs[0].runs[0] if cell1.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = (cell1.paragraphs[0].alignment if b != 3 else WD_ALIGN_PARAGRAPH.RIGHT)
                        val = df.iloc[a, b]
                        cell.text = (
                            "" if pd.isna(val) or val in [None, '']
                            else f'{val*100:,.1f}%' if isinstance(val, (float, int)) and b == 3 and table_to_replace.cell(a,b-2).text == 'Percentage'
                            else f'{val:,.1f}' if isinstance(val, (float, int)) and a in [5,6] and b == 3
                            else f'{val:,.0f}' if isinstance(val, (float, int))
                            else str(val)
                        )

                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[1]:
                cell1 = table_to_replace.cell(1,1)
                orig_run = cell1.paragraphs[0].runs[0] if cell1.paragraphs[0].runs else None
                font_name = orig_run.font.name if orig_run else None
                font_size = orig_run.font.size if orig_run else None
                font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                font_bold = orig_run.font.bold if orig_run else None
                
                for a in range(1,6):
                    for b in [2,3]:
                        alig = (cell1.paragraphs[0].alignment if b != 3 else WD_ALIGN_PARAGRAPH.RIGHT)
                        cell = table_to_replace.cell(a,b)
                        val = df.iloc[a,b]
                        cell.text = ('' if pd.isna(val) or val in [None, ""] else f'{val:,.0f}' if isinstance(val, (float, int)) else str(val))
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[2]:
                for a in range(1,30):
                    col0_val = df.iloc[a,0]
                    if pd.isna(col0_val) or col0_val in [None, '']:
                        for prev in range(a -1, -1, -1):
                            candidate = df.iloc[prev, 0]
                            if not pd.isna(candidate) and candidate not in [None, '']:
                                col0_val = candidate
                                break
                    for b in [2,3]:
                        alig = (cell1.paragraphs[0].alignment if b!=3 else WD_ALIGN_PARAGRAPH.RIGHT)
                        cell = table_to_replace.cell(a,b)
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val*100:,.0f}%' if isinstance(val, (float, int)) and b ==3 and table_to_replace.cell(a,b-2).text == 'Percentage'
                            else f'{val:,.0f}' if isinstance(val, (float, int)) and b == 3 and table_to_replace.cell(a,b-1).text == 'HART'
                            else f'${val:,.0f}' if isinstance(val, (float, int)) and 'income' in col0_val.lower()
                            else f'{val:,.1f}' if isinstance(val, (float, int))  and 'household size' in table_to_replace.cell(a,b-3).text.lower()
                            else f'{val:,.0f}' if isinstance(val, (float, int))
                            else str(val)
                            )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[3]:
                for a in range(1,30):
                    if a in list(range(2,12)):
                        cell = table_to_replace.cell(a,1)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,1]
                        cell.text = ('' if pd.isna(val) or val in [None, '']
                                    else str(val))
                        
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
                    for b in range(2,4):
                        alig = (cell1.paragraphs[0].alignment if b!=3 else WD_ALIGN_PARAGRAPH.RIGHT)
                        cell = table_to_replace.cell(a,b)
                        
                        
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val*100:,.1f}%' if  isinstance(val, (float, int)) and 'percentage' in table_to_replace.cell(a,b-2).text.lower()
                            else f'{val:,.0f}' if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[4]:
                for a in range(1, 31):
                    col0_val = df.iloc[a,0]
                    if pd.isna(col0_val) or col0_val in [None, '']:
                        for prev in range(a-1,-1,-1):
                            candidate = df.iloc[prev, 0]
                            if not pd.isna(candidate) and candidate not in [None, '']:
                                col0_val = candidate
                                break
                    for b in [2,3]:
                        cell = table_to_replace.cell(a,b)
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val*100:,.1f}%' if 'rate' in col0_val.lower() and isinstance(val, (float, int))
                            else f'{val:,.0f}' if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            alig = (cell1.paragraphs[0].alignment if b!=3 else WD_ALIGN_PARAGRAPH.RIGHT)
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[5]:
                for a in range(1,5):
                    for b in [2,3]:
                        cell = table_to_replace.cell(a,b)
                        val = df.iloc[a,b]
                        cell.text = ("" if pd.isna(val) or val in [None, ''] else f'{val:,.0f}' if isinstance(val, (float, int)) else str(val))
                        for paragraph in cell.paragraphs:
                            alig = (cell1.paragraphs[0].alignment if b!=3 else WD_ALIGN_PARAGRAPH.RIGHT)
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[6]:
                for a in range(1, 24):
                    for b in [2,3]:
                        cell = table_to_replace.cell(a,b)
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'${val:,.0f}' if isinstance(val, (float, int))
                            else str(val)
                            )
                        for paragraph in cell.paragraphs:
                            alig = (cell1.paragraphs[0].alignment if b!=3 else WD_ALIGN_PARAGRAPH.RIGHT)
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[7]:
                for a in range(1,16):
                    for b in [2,3]:
                        cell = table_to_replace.cell(a,b)
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val:,.0f}' if isinstance(val, (float, int))
                            else str(val)
                            )
                        for paragraph in cell.paragraphs:
                            alig = (cell1.paragraphs[0].alignment if b!=3 else WD_ALIGN_PARAGRAPH.RIGHT)
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[8]:
                for a in range(1,40):
                    for b in [2,3]:
                        cell = table_to_replace.cell(a,b)
                        val = df.iloc[a,b]
                        cell.text = (
                            "" if pd.isna(val) or val in [None, ""]
                            else f'{val*100:,.1f}%' if "percentage" in table_to_replace.cell(a,b-2).text.lower()  and isinstance(val,(float,int))
                            else f'{val:,.0f}' if isinstance(val,(float,int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            alig = (cell1.paragraphs[0].alignment if b!=3 else WD_ALIGN_PARAGRAPH.RIGHT)
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[9]:
                for a in range(1,6):
                    for b in range(1,5):
                        cell = table_to_replace.cell(a,b)
                        val = df.iloc[a,b]
                        cell.text = (
                            "" if pd.isna(val) or val in [None, ""]
                            else f'{val*100:,.0f}%' if b==4 and isinstance(val, (float, int))
                            else f'{val:,.0f}' if b == 3 and isinstance(val, (float, int))
                            else f'${val:,.0f}' if b ==2 and isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            alig = ( WD_ALIGN_PARAGRAPH.RIGHT)
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text ==titles[10]:
                for a in range(1,7):
                    for b in range(1,8):
                        cell = table_to_replace.cell(a,b)
                        val = df.iloc[a, b]
                        cell.text = (
                            "" if pd.isna(val) or val in [None, '']
                            else f'${val:,.0f}' if b ==1 and isinstance(val,(float,int))
                            else f'{val:,.0f}' if isinstance(val, (float,int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            alig = ( WD_ALIGN_PARAGRAPH.RIGHT)
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[11]:
                for a in range(1, 12):
                    for b in [2,3]:
                        cell = table_to_replace.cell(a,b)
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val*100:,.0f}%' if 'percentage' in table_to_replace.cell(a,b-2).text.lower() and isinstance(val, (float, int))
                            else f'{val:,.1f}' if isinstance(val, (float, int)) and a in [4,5]
                            else f'{val:,.0f}' if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            alig = (cell1.paragraphs[0].alignment if b!=3 else WD_ALIGN_PARAGRAPH.RIGHT)
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[12]:
                for a in range(1,27):
                    for b in [2,3]:
                        cell= table_to_replace.cell(a,b)
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val:,.1f}' if isinstance(val,(float, int)) and 'estimated' in table_to_replace.cell(a,b-1).text.lower()
                            else f'{val:,.0f}' if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            alig = (cell1.paragraphs[0].alignment if b!=3 else WD_ALIGN_PARAGRAPH.RIGHT)
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
            if search_text == titles[13]:
                for a in range(1, 7):
                    for b in range(1,7):
                        cell = table_to_replace.cell(a,b)
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else f'{val:,.0f}' if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            alig =  WD_ALIGN_PARAGRAPH.RIGHT
                            paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold

    #### Sheet: cover ####
    sheet = 'cover'
    columns = ['B:C']
    header_num = [1]
    rows_num = [7]

    titles = ['HOUSING NEEDS ASSESSMENT']

    for i in range(len(columns)):
        df = pd.read_excel(excel_file, sheet_name = sheet, usecols = columns[i], header = header_num[i], nrows=rows_num[i])
        search_text = titles[i]

        table_to_replace = None
        found_title = False
        body = doc.element.body
        for child in body:
            if not found_title:
                if child.tag ==qn('w:p'):
                    text = ''.join(t.text or '' for t in child.iter(qn('w:t')))
                    if search_text in text:
                        found_title = True
            else:
                if child.tag == qn('w:tbl'):
                    for tbl in doc.tables:
                        if tbl._element is child:
                            table_to_replace = tbl
                            break
                    break
        if table_to_replace and search_text == titles[0]:
                for a in range(0, 7):
                    for b in [1]:
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            '' if pd.isna(val) or val in [None, '']
                            else str(val.date()) if isinstance(val, (datetime.datetime, pd.Timestamp))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold

    #### Sheet: summary ####
    sheet = 'summary'
    columns = ['B:D', 'F:H', 'B:D', 'F:H', 'B:D']
    header_num = [3,3,16,16, 29]
    rows_num = [10,6,10,5, 8]
    titles=['DEMOGRAPHIC PROFILE', 'ECONOMIC PROFILE','HOUSING PROFILE', 'HOUSING NEED', 'UNIT DEMAND / PROJECTED NEEDS']

    for i in range(len(columns)):
        df = pd.read_excel(excel_file, sheet_name = sheet, usecols = columns[i], header = header_num[i], nrows = rows_num[i])
        search_text = titles[i]
        table_par_idx = None
        matches = [j for j, para in enumerate(doc.paragraphs) if search_text in para.text]
        table_par_idx = matches[2] if titles[i] in ['HOUSING NEED'] else matches[1] if len(matches) >1 and titles[i] in ['DEMOGRAPHIC PROFILE', 'ECONOMIC PROFILE','HOUSING PROFILE', 'UNIT DEMAND / PROJECTED NEEDS'] else matches[0] 
        if table_par_idx is not None:
            par_elem = doc.paragraphs[table_par_idx]._element
            sibling = par_elem.getnext()
            while sibling is not None:
                if sibling.tag == qn('w:tbl'):
                    for tbl in doc.tables:
                        if tbl._element is sibling:
                            table_to_replace = tbl
                            break
                    break
                sibling = sibling.getnext()
        if table_to_replace and search_text == titles[0]:
                for a in range(0, 10):
                    for b in range(0,3):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b]
                        cell.text = (
                            "" if pd.isna(val) or val in [None, '']
                            else f"{val:,.0f}" if isinstance(val, (float, int)) and (a,b) in [(0,0), (6,0), (8,2)]
                            else f'{val*100:,.0f}%' if isinstance(val, (float, int))
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
        if table_to_replace and search_text == titles[1]:
            for a in range(0,6):
                for b in range(0,3):
                    cell = table_to_replace.cell(a,b)
                    orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                    font_name = orig_run.font.name if orig_run else None
                    font_size = orig_run.font.size if orig_run else None
                    font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                    font_bold = orig_run.font.bold if orig_run else None
                    alig = cell.paragraphs[0].alignment
                    val = df.iloc[a,b]
                    cell.text = (
                        '' if pd.isna(val) or val in [None, '']
                        else f'${val:,.0f}' if isinstance(val, (float, int)) and (a,b) in [(0,0), (0,1)]
                        else f'{val:,.0f}' if isinstance(val, (float,int)) and (a,b) in [(2,0)]
                        else f'{val*100:,.0f}%' if isinstance(val, (float, int)) and (a,b) in [(4,1), (0,2), (4,0)]
                        else f'{val*100:,.1f}%' if isinstance(val, (float, int))
                        else str(val)
                    )
                    for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
        if table_to_replace and search_text == titles[2]:
            for a in range(0,10):
                for b in range(0,3):
                    cell = table_to_replace.cell(a,b)
                    orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                    font_name = orig_run.font.name if orig_run else None
                    font_size = orig_run.font.size if orig_run else None
                    font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                    font_bold = orig_run.font.bold if orig_run else None
                    alig = cell.paragraphs[0].alignment
                    val = df.iloc[a,b]
                    cell.text = (
                        '' if pd.isna(val) or val in [None, '']
                        else f'{val:,.0f}' if isinstance(val, (float, int)) and (a,b) in [(0,0),(2,0),(6,0),(8,1),(6,2),(8,2)]
                        else f'{val*100:,.1f}%' if isinstance(val, (float, int)) and (a,b) in [(4,1)]
                        else f'{val*100:,.0f}%' if isinstance(val, (float,int))
                        else str(val)
                    )
                    for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
        if table_to_replace and search_text == titles[3]:
            for a in range(0,5):
                for b in range(0,3):
                    cell = table_to_replace.cell(a,b)
                    orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                    font_name = orig_run.font.name if orig_run else None
                    font_size = orig_run.font.size if orig_run else None
                    font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                    font_bold = orig_run.font.bold if orig_run else None
                    alig = cell.paragraphs[0].alignment
                    val = df.iloc[a,b]
                    cell.text = (
                        '' if pd.isna(val) or val in [None, '']
                        else f'{val*100:,.0f}%' if isinstance(val, (float, int))
                        else str(val)
                    )
                    for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold
        if table_to_replace and search_text == titles[4]:
            for a in range(0,8):
                for b in range(0,3):
                    cell = table_to_replace.cell(a,b)
                    orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                    font_name = orig_run.font.name if orig_run else None
                    font_size = orig_run.font.size if orig_run else None
                    font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                    font_bold = orig_run.font.bold if orig_run else None
                    alig = cell.paragraphs[0].alignment
                    val = df.iloc[a,b]
                    cell.text = (
                        '' if pd.isna(val) or val in [None, '']
                        else f'{val:,.0f}' if isinstance(val, (float, int)) and (a,b) in [(0,0), (4,0), (6,0), (6,1),(2,2)]
                        else f'{val*100:,.0f}%' if isinstance(val, (float, int))
                        else str(val)
                    )
                    for paragraph in cell.paragraphs:
                        if alig is not None: paragraph.alignment = alig
                        for run in paragraph.runs:
                            if font_name: run.font.name = font_name
                            if font_size: run.font.size = font_size
                            if font_color: run.font.color.rgb = font_color
                            if font_bold is not None: run.font.bold = font_bold


    #### Sheet: summary ####
    sheet = 'priorities'
    columns = ['E:G']
    header_num = [2]
    rows_num = [6]
    titles=['PRIORITY ISSUES']

    for i in range(len(columns)):
        df = pd.read_excel(excel_file, sheet_name = sheet, usecols = columns[i], header = header_num[i], nrows = rows_num[i])
        search_text = titles[i]
        table_par_idx = None
        matches = [j for j, para in enumerate(doc.paragraphs) if search_text in para.text]
        table_par_idx = matches[2] if titles[i] in ['HOUSING NEED'] else matches[1] if len(matches) >1 else matches[0] 
        if table_par_idx is not None:
            par_elem = doc.paragraphs[table_par_idx]._element
            sibling = par_elem.getnext()
            while sibling is not None:
                if sibling.tag == qn('w:tbl'):
                    for tbl in doc.tables:
                        if tbl._element is sibling:
                            table_to_replace = tbl
                            break
                    break
                sibling = sibling.getnext()
        if table_to_replace and search_text == titles[0]:
                for a in range(1, 6):
                    for b in range(0,2):
                        cell = table_to_replace.cell(a,b)
                        orig_run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        font_name = orig_run.font.name if orig_run else None
                        font_size = orig_run.font.size if orig_run else None
                        font_color = orig_run.font.color.rgb if orig_run and orig_run.font.color.rgb else None
                        font_bold = orig_run.font.bold if orig_run else None
                        alig = cell.paragraphs[0].alignment
                        val = df.iloc[a,b+1]
                        cell.text = (
                            "" if pd.isna(val) or val in [None, '']
                            else str(val)
                        )
                        for paragraph in cell.paragraphs:
                            if alig is not None: paragraph.alignment = alig
                            for run in paragraph.runs:
                                if font_name: run.font.name = font_name
                                if font_size: run.font.size = font_size
                                if font_color: run.font.color.rgb = font_color
                                if font_bold is not None: run.font.bold = font_bold



    doc.save(temp_path)


    excel = win32com.client.Dispatch("Excel.Application")
    word = win32com.client.Dispatch("Word.Application")
    excel.Visible = False
    word.Visible = False

    wb = excel.Workbooks.Open(excel_path)
    doc = word.Documents.Open(temp_path)

    def paste_chart_at_caption(doc, chart_obj, caption_text):
        chart_obj.Copy()

        # default: append at end
        insert_pos = doc.Content.End - 1

        # find caption paragraph
        for para in doc.Paragraphs:
            if caption_text.lower() in para.Range.Text.lower():
                r = para.Range.Duplicate
                r.Collapse(0)  # end of paragraph
                insert_pos = r.Start
                break

        # delete existing inline shape at same position (reverse index loop is safer for COM)
        for idx in range(doc.InlineShapes.Count, 0, -1):
            shp = doc.InlineShapes(idx)
            if shp.Range.Start == insert_pos:
                shp.Delete()
                break

        # reacquire fresh range, then paste
        target = doc.Range(insert_pos, insert_pos)
        target.Paste()

    # ---- pop_now ----
    ws = wb.Sheets["popn_now"]
    pop_fig = [
        "Annual change in total population from previous year's population",
        "Annual regional demographic component shares of population change",
    ]
    for i, caption in enumerate(pop_fig, start=1):
        if i <= ws.ChartObjects().Count:
            paste_chart_at_caption(doc, ws.ChartObjects(i), caption)


    # ---- pop_fut ----

    ws = wb.Sheets['popn_fut']
    fig_text = ['Range of anticipated total population over next 10 years']

    for i, caption in enumerate(fig_text, start=1):
        if i <= ws.ChartObjects().Count:
            paste_chart_at_caption(doc, ws.ChartObjects(i), caption)

    # ---- cal ----

    ws = wb.Sheets['projection']
    fig_text = ['Anticipated running dwelling shortage',
                'Anticipated new dwelling demand by dwelling type, over next 10 years',
                'Anticipated new dwelling demand by dwelling price level, over next 10 years',
                'Impact of regionally influenced demand on anticipated overall demand over next 10 years']
    chart_num = [9, 7, 8, 10]

    for caption, chart_idx in zip(fig_text, chart_num):
        if chart_idx <= ws.ChartObjects().Count:
            paste_chart_at_caption(doc, ws.ChartObjects(chart_idx), caption)




    base_name = os.path.splitext(filename)[0]
    excel_number = ''.join(ch for ch in base_name if ch.isdigit())
    output_word_path = os.path.join(output_folder, f"HNA_MUNICIPAL_{excel_number}.docx")
    doc.SaveAs(output_word_path)
    doc.Close()
    wb.Close()
    excel.Quit()
    word.Quit()
    os.remove(temp_path)


